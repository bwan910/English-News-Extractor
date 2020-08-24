import nltk
from newspaper import Article
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import requests  # to get image from the web
import shutil  # to save it locally
from tkinter import *
from tkinter import messagebox


def main():
    def newspaper():
        document = Document()

        url = url_input.get()

        article = Article(url)

        article.download()
        article.parse()

        # adding a heading
        title = document.add_heading(article.title)
        # format the title to center
        title_format = title.paragraph_format
        title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # adding a publish date
        document.add_heading("Publish Date: " +
                             str(article.publish_date), level=3)

        # needed to be convert to a string in able to write to document
        listToStr = ' '.join([str(elem) for elem in article.authors])

        document.add_heading("Authors: " + listToStr + '\n', level=4)

        # Set up the image URL and filename
        image_url = article.top_image
        filename = image_url.split("/")[-1]

        # Open the url image, set stream to True, this will return the stream content.
        # Use stream = True to guarantee no interruptions.
        r = requests.get(image_url, stream=True)

        # Check if the image was retrieved successfully
        if r.status_code == 200:
            # Set decode_content value to True, otherwise the downloaded image file's size will be zero.
            r.raw.decode_content = True

            # Open a local file with wb ( write binary ) permission.
            with open(filename, 'wb') as f:
                shutil.copyfileobj(r.raw, f)

            print('Image sucessfully Downloaded: ', filename)
        else:
            print('Image Couldn\'t be retreived')

        # adding the top image
        document.add_picture(filename, width=Inches(5.00))

        # adding the main content
        document.add_paragraph(article.text)

        # save document
        urlSplit = url.split("/")[-1:]
        urlToString = ' '.join([str(elem) for elem in urlSplit])
        document.save(filename_input.get()+'.docx')

        # display succesful message box
        url_input.delete(0, "end")
        filename_input.delete(0, "end")
        messagebox.showinfo("Successfully", "Document saved")

    root = Tk()
    root.geometry('500x500')
    root.title("Save news to word documents from URL ")

    # Label for url
    url_label = Label(root, text='URL', font=('Arial', 16))
    url_label.place(x=50, y=100)

    # Input box for url
    url_input = Entry(root, bd=2, font=('Arial', 12), width=30)
    url_input.place(x=150, y=100)

    # Label for filename
    filename_label = Label(root, text='Filename', font=('Arial', 16))
    filename_label.place(x=50, y=150)

    # Input box for filename
    filename_input = Entry(root, bd=2, font=('Arial', 12), width=30)
    filename_input.place(x=150, y=150)

    # save button
    btn_save = Button(root, text="Save", font=(
        'Bold', 18), bg="Yellow", command=newspaper)
    btn_save.place(x=200, y=250)

    root.mainloop()


main()
